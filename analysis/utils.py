from django.contrib.auth.models import Group
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import csv
from docx import Document
import PyPDF2

ADMIN_GROUP = "Admin"
TEACHER_GROUP = "Teacher"

def ensure_core_groups():
    for name in [ADMIN_GROUP, TEACHER_GROUP]:
        Group.objects.get_or_create(name=name)
        
from docx import Document

def read_word_answer_key(file):
    """
    Reads a Word (.docx) answer key file and returns a list of answers.
    Each line in the file should be like "1. A", "2. B", etc.
    """
    answers = []
    # If file is InMemoryUploadedFile, you can use file.read() as bytes
    if hasattr(file, 'read'):
        doc = Document(file)
    else:
        doc = Document(file)

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        parts = line.split('.')
        if len(parts) >= 2:
            answer = parts[1].strip()
            if answer in ["A", "B", "C", "D", "E"]:
                answers.append(answer)
    return answers

def generate_bubble_sheet_pdf(answers, pdf_path):
    """
    Generate a PDF bubble sheet from a list of answers.
    answers: ['A', 'B', 'C', ...]
    pdf_path: full path to save PDF
    """
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    x_start = 100
    y_start = height - 100
    bubble_size = 15
    gap_x = 50
    gap_y = 30
    options = ["A", "B", "C", "D", "E"]

    # Draw header
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x_start, y_start + 20, "Bubble Test Sheet")

    # Draw bubbles
    for i, ans in enumerate(answers):
        y = y_start - (i * gap_y)
        c.drawString(x_start - 30, y, str(i+1))
        for j, option in enumerate(options):
            x = x_start + j * gap_x
            c.circle(x, y, bubble_size/2)
            if option == ans:
                c.setFillColorRGB(0, 0, 0)
                c.circle(x, y, bubble_size/2, fill=1)
                c.setFillColorRGB(0, 0, 0)  # Reset
    c.save()

def read_answer_key_file(file):
    answers = []
    ext = getattr(file, "name", "").split(".")[-1].lower()

    try:
        # ---------- DOC / DOCX ----------
        if ext in ["docx", "doc"]:
            doc = Document(file)
            for para in doc.paragraphs:
                text = para.text.strip()
                parts = text.split()
                if parts:
                    last = parts[-1].upper()
                    if last in ["A", "B", "C", "D", "E"]:
                        answers.append(last)

        # ---------- CSV ----------
        elif ext == "csv":
            decoded = file.read().decode("utf-8").splitlines()
            reader = csv.reader(decoded)
            for row in reader:
                if row:
                    val = row[-1].strip().upper()
                    if val in ["A", "B", "C", "D", "E"]:
                        answers.append(val)

        # ---------- PDF ----------
        elif ext == "pdf":
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text = page.extract_text()
                if not text:
                    continue

                for raw in text.splitlines():
                    line = raw.strip()
                    if not line:
                        continue

                    cleaned = (
                        line.replace(")", "")
                            .replace(".", "")
                            .replace("-", "")
                            .replace(":", "")
                            .strip()
                    )
                    parts = cleaned.split()
                    if not parts:
                        continue

                    last = parts[-1].upper()
                    if last in ["A", "B", "C", "D", "E"]:
                        answers.append(last)

        else:
            print("Unsupported file type:", ext)

    except Exception as e:
        print("Error reading file:", e)

    return answers


from .models import Institution, InstitutionAdmin, Teacher

def get_current_institution(request):
    # ✅ avoid AnonymousUser crash
    if not getattr(request, "user", None) or not request.user.is_authenticated:
        return None

    # ✅ if admin account mapped to an institution
    inst_admin = (
        InstitutionAdmin.objects
        .select_related("institution")
        .filter(user=request.user)
        .first()
    )
    if inst_admin and inst_admin.institution:
        return inst_admin.institution

    # ✅ if teacher account belongs to an institution
    teacher = (
        Teacher.objects
        .select_related("institution")
        .filter(user=request.user)
        .first()
    )
    if teacher and teacher.institution:
        return teacher.institution

    # ✅ fallback: session (optional)
    inst_id = request.session.get("institution_id")
    if inst_id:
        return Institution.objects.filter(id=inst_id).first()

    return None

def get_interpretation(status):
    mapping = {
        'correct': 'Correct response.',
        'incorrect': 'Incorrect response.',
        'blank': 'No mark detected.',
        'invalid': 'Messy or multiple marks detected.'
    }
    return mapping.get(status, 'Unknown status.')